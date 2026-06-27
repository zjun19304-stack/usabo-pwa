#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成 USABO 题目整理模板文件 (questions_template.docx)
运行：python make_template.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页面设置（A4）────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 辅助函数 ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1B, 0x6C, 0xA8)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_question_table(doc, fields):
    """
    用表格形式添加一道题
    fields: list of (label, value, label_bg) tuples
    """
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    for label, value, *extras in fields:
        label_bg = extras[0] if extras else 'E8F0FA'
        row = table.add_row()
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(13)

        # 标签列
        c0 = row.cells[0]
        c0.text = label
        if c0.paragraphs[0].runs:
            c0.paragraphs[0].runs[0].bold = True
            c0.paragraphs[0].runs[0].font.size = Pt(10)
            c0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1B, 0x6C, 0xA8)
        set_cell_bg(c0, label_bg)

        # 值列
        c1 = row.cells[1]
        c1.text = value
        if c1.paragraphs[0].runs:
            c1.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(c1, 'FFFFFF')

    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
#  封面标题
# ════════════════════════════════════════════════════════════
title = doc.add_heading('USABO 题目整理模板', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1B, 0x6C, 0xA8)

sub = doc.add_paragraph('美国生物奥林匹克竞赛 · 题库录入专用模板')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
sub.runs[0].font.size = Pt(11)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
#  字段说明表（简化版）
# ════════════════════════════════════════════════════════════
add_section_title(doc, '📋 一、字段说明')

info_table = doc.add_table(rows=1, cols=3)
info_table.style = 'Table Grid'
hdr = info_table.rows[0]
for i, txt in enumerate(['字段', '必填', '说明']):
    hdr.cells[i].text = txt
    if hdr.cells[i].paragraphs[0].runs:
        hdr.cells[i].paragraphs[0].runs[0].bold = True
        hdr.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr.cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(hdr.cells[i], '1B6CA8')

field_rows = [
    ('【板块】', '✅ 必填', 'ECO / ANB / GEN / HER / RPH / GTE / BCC / APH / PLA / TAX / VIR / MIT'),
    ('【难度】', '可选',   '1=简单 / 2=中等 / 3=困难，默认 2（可不填）'),
    ('【题干】', '✅ 必填', '题目正文，可以多行'),
    ('【图片】', '可选',   '有图片时填写路径如 images/eco_001.png，无图片直接删掉此行'),
    ('A. B. C.', '✅ 必填', '选项内容，每行一个，至少 2 个'),
    ('【答案】', '✅ 必填', '单选填字母如 B；多选用逗号如 A,C；T/F判断题见下方说明'),
    ('【解析】', '可选',   '答案解析，建议填写'),
    ('【来源】', '可选',   '如 USABO Open Exam 2022'),
    ('---',      '✅ 必填', '每道题结束后必须有这一行作为分隔符'),
]

for f, req, desc in field_rows:
    row = info_table.add_row()
    row.cells[0].text = f
    row.cells[1].text = req
    row.cells[2].text = desc
    for cell in row.cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# T/F 说明
p_tf = doc.add_paragraph()
r_tf = p_tf.add_run('💡 T/F判断题说明：')
r_tf.bold = True
r_tf.font.size = Pt(10)
r_tf.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)
doc.add_paragraph(
    '选项直接写每个陈述的内容（如 A. 细胞膜具有选择透过性），答案写每项对应的 T 或 F，用逗号隔开。\n'
    '例：4个选项答案为 "T,F,T,F" 表示 A=True，B=False，C=True，D=False。\n'
    '程序会根据答案格式自动识别题型，不需要手动指定。'
).runs[0].font.size = Pt(9)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════
#  示例题
# ════════════════════════════════════════════════════════════
add_section_title(doc, '📝 二、示例题（请参考此格式填写）')

# ── 示例 1：单选题
p1 = doc.add_paragraph()
p1.add_run('示例 1 — 单选题').bold = True

add_question_table(doc, [
    ('【板块】', 'ECO'),
    ('【难度】', '2'),
    ('【题干】', 'Which of the following best describes the role of decomposers in an ecosystem?'),
    ('A.',      'They convert solar energy into chemical energy'),
    ('B.',      'They break down organic matter and return nutrients to the soil'),
    ('C.',      'They occupy the highest trophic level in a food chain'),
    ('D.',      'They compete with producers for available light'),
    ('【答案】', 'B'),
    ('【解析】', 'Decomposers (bacteria and fungi) break down dead organic matter, releasing nutrients back into the environment for producers to use.'),
    ('【来源】', 'USABO Open Exam 2022'),
    ('---',     ''),
])

# ── 示例 2：多选题
p2 = doc.add_paragraph()
p2.add_run('示例 2 — 多选题（答案含多个字母）').bold = True

add_question_table(doc, [
    ('【板块】', 'GEN'),
    ('【难度】', '3'),
    ('【题干】', "Which of the following are true regarding DNA replication in eukaryotes? (Select all that apply)"),
    ('A.',      "Replication occurs in the 5' to 3' direction"),
    ('B.',      'Okazaki fragments are formed on the leading strand'),
    ('C.',      'Multiple origins of replication are used simultaneously'),
    ('D.',      'DNA polymerase can initiate new strands without a primer'),
    ('【答案】', 'A,C'),
    ('【解析】', "DNA is synthesized 5'→3'. Okazaki fragments are on the lagging strand. Eukaryotes use multiple origins. DNA polymerase requires a primer."),
    ('【来源】', 'USABO Invitational 2023'),
    ('---',     ''),
])

# ── 示例 3：T/F 判断题
p3 = doc.add_paragraph()
p3.add_run('示例 3 — T/F 判断题（答案写 T 或 F，用逗号隔开）').bold = True

add_question_table(doc, [
    ('【板块】', 'BCC'),
    ('【难度】', '2'),
    ('【题干】', 'For each of the following statements about enzyme function, indicate whether it is True (T) or False (F).'),
    ('A.',      'Enzymes lower the activation energy of a reaction'),
    ('B.',      'Enzymes are consumed during the reaction they catalyze'),
    ('C.',      'Enzyme activity can be affected by changes in pH'),
    ('D.',      'A competitive inhibitor permanently binds to the active site'),
    ('【答案】', 'T,F,T,F'),
    ('【解析】', 'Enzymes lower activation energy (T). Enzymes are NOT consumed—they are recycled (F). pH affects enzyme shape and activity (T). Competitive inhibition is reversible (F).'),
    ('【来源】', 'USABO Open Exam 2021'),
    ('---',     ''),
])

# ── 示例 4：有图片的题目
p4 = doc.add_paragraph()
p4.add_run('示例 4 — 有图片的题目（填写【图片】字段）').bold = True

add_question_table(doc, [
    ('【板块】', 'MIT'),
    ('【难度】', '3'),
    ('【题干】', 'Based on the diagram shown, which stage of mitosis is depicted?'),
    ('【图片】', 'images/mit_001.png'),
    ('A.',      'Prophase'),
    ('B.',      'Metaphase'),
    ('C.',      'Anaphase'),
    ('D.',      'Telophase'),
    ('【答案】', 'B'),
    ('【解析】', 'In metaphase, chromosomes align along the cell equator (metaphase plate). This is clearly shown in the diagram with chromosomes lined up in the middle.'),
    ('【来源】', 'USABO Semifinal 2023'),
    ('---',     ''),
])

# ════════════════════════════════════════════════════════════
#  空白模板
# ════════════════════════════════════════════════════════════
add_section_title(doc, '✏️ 三、在这里填写你的题目（复制空白模板，按需增加）')

note = doc.add_paragraph(
    '提示：复制下方表格，填入你的题目。有图片则保留【图片】行，没有则删掉。'
    '题型不需要手动指定，程序根据答案自动判断。'
)
note.runs[0].font.size = Pt(9)
note.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

for i in range(1, 6):
    p_i = doc.add_paragraph()
    r_i = p_i.add_run(f'题目 {i}')
    r_i.bold = True
    r_i.font.size = Pt(10)
    r_i.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)

    add_question_table(doc, [
        ('【板块】', ''),
        ('【难度】', '2'),
        ('【题干】', ''),
        ('【图片】', '（无图片则删掉此行）'),
        ('A.',      ''),
        ('B.',      ''),
        ('C.',      ''),
        ('D.',      ''),
        ('【答案】', ''),
        ('【解析】', ''),
        ('【来源】', ''),
        ('---',     ''),
    ])

# ════════════════════════════════════════════════════════════
#  导入命令
# ════════════════════════════════════════════════════════════
add_section_title(doc, '🚀 四、填完后如何导入')

cmds = [
    '1. 保存这个文件（另存为 .docx 格式）',
    '2. 在终端运行（追加模式，不影响已有题目）：',
    '   python doc_to_questions.py  你的文件.docx  --out ../questions.js  --append',
    '',
    '3. 重新加密（把学生密码填上）：',
    '   python encrypt_questions.py  "学生A:密码A,学生B:密码B"',
    '',
    '4. 推送到 GitHub：',
    '   git add questions.enc.js && git commit -m "Add questions" && git push',
]
for cmd in cmds:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = 'Courier New' if cmd.startswith('   ') else None
    run.font.size = Pt(9)

# ── 保存 ──────────────────────────────────────────────────────────────────────
out_path = 'questions_template.docx'
doc.save(out_path)
print(f"✅ 模板已生成：{out_path}")
